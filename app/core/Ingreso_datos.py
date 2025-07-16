from docx import Document
import os
import config.rutas as dir
from datetime import datetime
def generar_presupuesto_docx(lista_items, nombre_cliente):
    import os
    import config.rutas as dir
    from docx import Document

    base_dir = os.path.dirname(dir.BASE_PATH)
    ruta_plantilla = os.path.join(base_dir, "app", "resources", "Plantillas", "presupuesto_coctel.docx")
    doc = Document(ruta_plantilla)

    # Obtener fecha y hora en formato seguro para nombres de archivo
    ahora = datetime.now()
    fecha_str = ahora.strftime("%d-%m-%Y")
    hora_str = ahora.strftime("%H-%M")  # Usar guiones para evitar dos puntos

    # Reemplazar marcadores dentro del Word (si existen)
    for para in doc.paragraphs:
        if "{{cliente}}" in para.text:
            para.text = para.text.replace("{{cliente}}", nombre_cliente)
        if "{{fecha}}" in para.text:
            para.text = para.text.replace("{{fecha}}", f"{fecha_str} {hora_str}")

    # Agregar la tabla con productos (como ya haces normalmente)
    total_general = 0
    table = next((t for t in doc.tables if "Descripción" in t.rows[0].cells[1].text), None)
    if not table:
        raise Exception("No se encontró una tabla válida en la plantilla.")

    for cantidad, descripcion, total in lista_items:
        row = table.add_row().cells
        row[0].text = str(cantidad)
        row[1].text = descripcion
        row[2].text = f"${total:,}".replace(",", ".")
        total_general += total

    for para in doc.paragraphs:
        if "Total:" in para.text:
            para.text = f"Total: ${total_general:,}".replace(",", ".")
            break

    # Crear nombre del archivo con formato correcto
    nombre_archivo = f"Cotizacion_{nombre_cliente.replace(' ', '_')}_{fecha_str}_{hora_str}.docx"
    ruta_salida = os.path.join(base_dir, "app", "resources", "Salida", nombre_archivo)

    doc.save(ruta_salida)
    return ruta_salida